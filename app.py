import streamlit as st
import pandas as pd
from io import BytesIO
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Cm, Pt, RGBColor
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn
import re

# إعدادات واجهة المستخدم
st.set_page_config(page_title="نظام تنسيق وتدقيق كشوفات الوكلاء", layout="wide")
st.markdown("""
    <style>
    th, td { text-align: right !important; dir: rtl !important; }
    div.stButton > button { background-color: #2E4053; color: white; width: 100%; font-weight: bold; border-radius: 8px; font-size: 18px;}
    .report-box { background-color: #F4F6F7; padding: 15px; border-radius: 8px; border-right: 5px solid #2E4053; text-align: right; margin-bottom: 10px;}
    </style>
""", unsafe_allow_html=True)

st.markdown("<h1 style='text-align: right;'>نظام تنسيق وتدقيق كشوفات الوكلاء المطور 📄💎</h1>", unsafe_allow_html=True)

if "processing_done" not in st.session_state:
    st.session_state.processing_done = False
    st.session_state.df_final = None
    st.session_state.output_filename = ""

# -----------------------------------------------------------------------------
# مساعدات التنسيق المتقدمة لملفات Word عبر الـ XML
# -----------------------------------------------------------------------------
def set_table_borders(table, color_hex="2A4B7C"):
    tblPr = table._tbl.tblPr
    borders = parse_xml(f'''
        <w:tblBorders {nsdecls("w")}>
            <w:top w:val="single" w:sz="6" w:space="0" w:color="{color_hex}"/>
            <w:left w:val="single" w:sz="6" w:space="0" w:color="{color_hex}"/>
            <w:bottom w:val="single" w:sz="6" w:space="0" w:color="{color_hex}"/>
            <w:right w:val="single" w:sz="6" w:space="0" w:color="{color_hex}"/>
            <w:insideH w:val="single" w:sz="4" w:space="0" w:color="{color_hex}"/>
            <w:insideV w:val="single" w:sz="4" w:space="0" w:color="{color_hex}"/>
        </w:tblBorders>
    ''')
    tblPr.append(borders)

def set_cell_background(cell, fill_hex):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_no_wrap(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    no_wrap = parse_xml(f'<w:noWrap {nsdecls("w")}/>')
    tcPr.append(no_wrap)

def format_cell_advanced(cell, text, bold=False, color_rgb=None, size_pt=14, font_name="Calibri", align="center"):
    cell.text = str(text)
    p = cell.paragraphs[0]
    
    if align == "right":
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif align == "left":
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
    pPr = p.paragraph_format.element.get_or_add_pPr()
    pPr.append(parse_xml(f'<w:bidi {nsdecls("w")}/>'))
    
    for run in p.runs:
        run.bold = bold
        if color_rgb:
            run.font.color.rgb = color_rgb
            
        rPr = run._r.get_or_add_rPr()
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), font_name)
        rFonts.set(qn('w:hAnsi'), font_name)
        rFonts.set(qn('w:cs'), font_name)
        rPr.append(rFonts)
        run.font.size = Pt(size_pt)

# -----------------------------------------------------------------------------
# محرك قراءة وتنظيف البيانات
# -----------------------------------------------------------------------------
def extract_and_clean_data(file_obj):
    doc = Document(file_obj)
    raw_records = []
    
    # عداد لتسجيل ترتيب الاسم في الملف الأصلي
    original_seq_counter = 1 
    
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
            if not any(cells) or "المركز" in "".join(cells) or "الوكيل" in "".join(cells) or "اسم رب" in "".join(cells):
                continue
            
            # إيجاد حقل الاسم
            name_idx = -1
            max_len = 0
            for i, c in enumerate(cells):
                if any('\u0600' <= char <= '\u06FF' for char in c) and not any(char.isdigit() for char in c):
                    if len(c) > max_len:
                        max_len = len(c)
                        name_idx = i
            if name_idx == -1: continue
            
            # إيجاد رقم البطاقة القديم
            card_indices = [i for i, c in enumerate(cells) if c.isdigit() and len(c) >= 5]
            if not card_indices: continue
            old_card_num = cells[card_indices[0]]
                
            raw_records.append({
                "التسلسل الأصلي": str(original_seq_counter),
                "اسم رب الأسرة": cells[name_idx],
                "رقم البطاقة القديم": old_card_num
            })
            
            # زيادة العداد لكل قيد صحيح يتم قراءته
            original_seq_counter += 1
            
    df = pd.DataFrame(raw_records)
    if not df.empty:
        df = df.sort_values(by="اسم رب الأسرة").reset_index(drop=True)
        df.insert(0, "ت", df.index + 1)
    return df

# -----------------------------------------------------------------------------
# محرك بناء تقرير Word الاحترافي الجديد
# -----------------------------------------------------------------------------
def build_professional_word_report(df, filename_base):
    doc = Document()
    
    for section in doc.sections:
        section.top_margin = Cm(0.5)
        section.bottom_margin = Cm(0.5)
        section.left_margin = Cm(0.3)
        section.right_margin = Cm(0.3)
        
    clean_name = filename_base
    words_to_remove = ["مستكشف", "معدل", "كشف", "منسق", "جاهز"]
    for w in words_to_remove:
        clean_name = clean_name.replace(w, "")
    clean_name = re.sub(r'[a-zA-Z]', '', clean_name)
    clean_name = re.sub(r'[\-_+_.]', '', clean_name)
    clean_name = " ".join(clean_name.split())
    
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    title_run = title_p.add_run(f"الكشف الإحصائي المنسق للوكيل: {clean_name}")
    title_run.font.name = "Segoe UI Semibold"
    title_run.font.size = Pt(14)
    title_run.bold = True
    
    footer = doc.sections[0].footer
    footer_p = footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    f_run = footer_p.add_run("صفحة ")
    f_run.font.size = Pt(10)
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
    fldChar3 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    f_run._r.extend([fldChar1, instrText, fldChar2, fldChar3])
    
    headers = ["ت", "اسم رب الأسرة", "رقم البطاقة القديم", "الأصلي"] + [str(i) for i in range(1, 13)]
    num_cols = len(headers) 
    
    table = doc.add_table(rows=1, cols=num_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    set_table_borders(table, color_hex="2A4B7C")
    
    tblPr = table._tbl.tblPr
    tblPr.append(parse_xml(f'<w:bidiVisual {nsdecls("w")}/>'))
    
    trPr = table.rows[0]._tr.get_or_add_trPr()
    trPr.append(parse_xml(f'<w:tblHeader {nsdecls("w")}/>'))
    
    # 🔴 التعديلات على مقاسات الأعمدة 🔴
    col_widths = [
        Cm(0.8),              # ت
        Cm(7.5),              # اسم رب الأسرة (حجم ثابت 7.5 سم)
        Cm(2.6),              # رقم البطاقة القديم
        Cm(1.2),              # التسلسل الأصلي
    ] + [Cm(1.5)] * 12        # 12 عمود بحجم مضاعف (1.5 سم)
    
    COLOR_NAVY_BLUE = RGBColor(42, 75, 124)
    
    for i, title in enumerate(headers):
        hdr_cells = table.rows[0].cells
        hdr_cells[i].width = col_widths[i]
        
        cell_align = "center"
        font_size = 11 if i >= 4 else 12 
        format_cell_advanced(hdr_cells[i], title, bold=True, size_pt=font_size, font_name="Segoe UI Semibold", align=cell_align, color_rgb=COLOR_NAVY_BLUE)
            
    HEX_LIGHT_GREY = "F2F4F4"   
    HEX_FAINT_SKY = "EAF2F8"    
    
    for idx, row in df.iterrows():
        row_cells = table.add_row().cells
        
        r_trPr = table.rows[idx+1]._tr.get_or_add_trPr()
        r_trPr.append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))
        
        for i in range(num_cols):
            row_cells[i].width = col_widths[i]
            
        set_cell_no_wrap(row_cells[1])
        
        for i in range(num_cols):
            val = ""
            cell_align = "center"
            font_size = 14 
            
            if i == 0: val = row["ت"]
            elif i == 1: 
                val = row["اسم رب الأسرة"]
                cell_align = "left" 
                font_size = 16 
            elif i == 2: val = row["رقم البطاقة القديم"]
            elif i == 3: val = row["التسلسل الأصلي"]
            
            format_cell_advanced(row_cells[i], val, size_pt=font_size, font_name="Calibri", align=cell_align)
            
            if i == 0: set_cell_background(row_cells[i], HEX_LIGHT_GREY)
            elif i == 2: set_cell_background(row_cells[i], HEX_FAINT_SKY)

    total_all = len(df)
    
    doc.add_paragraph() 
    stats_p = doc.add_paragraph()
    stats_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    stats_p.paragraph_format.element.get_or_add_pPr().append(parse_xml(f'<w:bidi {nsdecls("w")}/>'))
    
    stats_text = f"العدد الكلي للأسر = {total_all}"
    stats_run = stats_p.add_run(stats_text)
    stats_run.font.name = "Segoe UI Semibold"
    stats_run.font.size = Pt(13)
    stats_run.bold = True
    stats_run.font.color.rgb = COLOR_NAVY_BLUE

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# -----------------------------------------------------------------------------
# واجهة استخدام التطبيق
# -----------------------------------------------------------------------------
st.markdown("<h3 style='text-align: right;'>📂 رفع الكشف المراد تدقيقه وتنسيقه للمطبعة</h3>", unsafe_allow_html=True)
uploaded_file = st.file_uploader("ارفع كشف الوكلاء", type=['docx'], key="doc_input_v8", label_visibility="collapsed")

st.markdown("<br>", unsafe_allow_html=True)

if uploaded_file:
    current_filename = uploaded_file.name.rsplit('.', 1)[0]
    if st.session_state.output_filename != current_filename:
        st.session_state.processing_done = False

if st.button("⚙️ تشغيل محرك التنظيم والتنسيق المتقدم الكلي"):
    if uploaded_file:
        with st.spinner('جاري الترتيب وضبط المقاسات...'):
            try:
                df_res = extract_and_clean_data(uploaded_file)
                if not df_res.empty:
                    st.session_state.df_final = df_res
                    st.session_state.output_filename = uploaded_file.name.rsplit('.', 1)[0]
                    st.session_state.processing_done = True
                else:
                    st.error("لم يتم العثور على بيانات جداول متوافقة.")
            except Exception as e:
                st.error(f"خطأ غير متوقع: {e}")
    else:
        st.warning("الرجاء رفع ملف docx أولاً.")

if st.session_state.processing_done:
    df_final = st.session_state.df_final
    output_filename = st.session_state.output_filename
    
    st.success(f"✅ تم التنظيم الأبجدي بنجاح لـ ({len(df_final)}) قيد اسم.")
    
    with st.spinner('جاري صياغة وهيكلة مستند Word المطور...'):
        word_output = build_professional_word_report(df_final, output_filename)
        
    st.download_button(
        label="📥 تحميل كشف الوكلاء المنسق والجاهز (Word)",
        data=word_output,
        file_name=f"كشف_منسق_جاهز_{output_filename}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
